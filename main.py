#!/usr/bin/env python
# -*- coding: utf-8 -*-
# @Time : 2019/11/20 14:18
# @Author : Qi Meng
# @File : main.py

import os
import re
import copy
import time
from tika import parser
from docx import Document
from html import unescape
from configparser import ConfigParser


def save_text(content1):
    with open(text_file, "w", encoding='utf-8') as f:
        f.write(str(content1))
        f.close()


def save_text_to_word(content1, file_path):
    doc = Document()
    for line in content1.split('\n'):
        paragraph = doc.add_paragraph()
        paragraph.add_run(remove_control_characters(line))
    doc.save(file_path)


def remove_control_characters(content2):
    mpa = dict.fromkeys(range(32))
    return content2.translate(mpa)


def html_to_plain_text ( html ):
# 处理html
    text = re.sub('<head.*?>.*?</head>', '', html, flags=re.M | re.S | re.I)
    text = re.sub(r'<a\s.*?>', ' HYPERLINK ', text, flags=re.M | re.S | re.I)
    text = re.sub('<.*?>', '', text, flags=re.M | re.S)
    text = re.sub(r'-$\n', '', text, flags=re.M | re.S)
    text = re.sub(r'\n(\n+)', '---<<<///qm', text, flags=re.M | re.S)
    text = re.sub(r'(\s*\n)+', '', text, flags=re.M | re.S)
    text = re.sub('---<<<///qm', '\n', text, flags=re.M | re.S)
# 处理特殊行
    text = re.sub('^[A-Z|\\s|0-9]+$\n', '\n', text, flags=re.M | re.S)
# 处理换行问题
    text = re.sub(r"([A-Z][a-zA-Z]+)\s*$\n", r"\1!qm! \n", text, flags=re.M | re.S)
    text = re.sub(r"(\.|\?|\!|。|？|！)$\n", r"\1 \n", text, flags=re.M | re.S)
    text = re.sub(r"(?<!((\.|\?|\!|。|？|！)\s))$\n", "", text, flags=re.M | re.S)
    text = re.sub(r"!qm!", r"", text, flags=re.M | re.S)
# 处理空格问题
    text = re.sub(r'\s([a-zA-Z])\s([a-zA-Z])\s', r'\1\2', text, flags=re.M | re.S)
    return unescape(text)


config_parser = ConfigParser()
config_parser.read('config.cfg', encoding='utf-8')
config = config_parser['default']
spy_folder = config['target_folder']
key_word = config['key_word'].split("\"")[1]
print("关键词为：", key_word)

flag = input("是否需要重新指定关键词？\n\t1：需要;\n\t2：不需要。\n")
if flag == '1':
    key_word = input("请输入关键词：\n")
print("\n关键词为：", key_word)

flag = input("请选择处理方式：\n\t1：browsing;\n\t2：batch.\n")
if flag == '2':
    list1 = []
    print("开始批量处理文件夹：", spy_folder)
else:
    list1 = os.listdir(spy_folder)
    print("开始监听文件夹：", spy_folder)

while 1:
    list2 = os.listdir(spy_folder)         # 检测文件夹是否更新
    if list1 == list2:                     # 如果未更新就休息之后继续监听
        time.sleep(1)
        continue

    new_file = list(set(list2).difference(set(list1)))   # 如果更新了，new_file为新增文件的list
    list1 = copy.deepcopy(list2)                         # list1 更新

    for file in new_file:                                # 依次处理新增的每个文件
        print("\n新增文件：", file)
        extension_name = os.path.splitext(file)[1]       # 获取文件后缀对文件进行判断
        if extension_name != '.pdf':
            print("非pdf文件！")
            continue
        file_name = os.path.splitext(file)[0]            # 获取文件名
        print("正在处理：", file_name, "请稍等。。。")

        pdf_file = spy_folder + '/' + file     # pdf_file 为源文件的完整地址
        word_file = config['save_folder'] + '/' + file_name + '.docx'     # word_file 为文件存储地址
        # text_file = config['save_folder'] + '/' + file_name + '.txt'     # word_file 为文件存储地址
        parse_entire_pdf = parser.from_file(pdf_file, xmlContent=True)
        parse_entire_pdf = parse_entire_pdf['content']
        content = html_to_plain_text(parse_entire_pdf)

        pattern = re.compile(".*" + key_word + ".*", re.I)
        result = pattern.findall(content)
        print("共计段数：", len(result))
        result_text = ""
        cnt = 1
        for para in result:
            # para = re.sub(key_word, "  ——>" + key_word + "<——  ", para, flags = re.I)
            print("第", cnt, "段：", para, "\n")
            result_text += para
            result_text += "\n"
            cnt += 1
        # save_text(result_text)
        save_text_to_word(result_text, word_file)
        print("处理完毕！")
    if flag == '2':
        break
os.system("pause")

